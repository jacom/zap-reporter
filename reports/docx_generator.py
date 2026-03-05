"""Agreement document generator — docxtpl + LibreOffice headless.

Flow:
  PentestAgreement  →  docxtpl render  →  .docx  →  (LibreOffice)  →  .pdf
"""

import io
import logging
import os
import subprocess
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'agreement_template.docx')
FONT          = 'Sarabun'


# ─────────────────────────────────────────────────────────────────────────────
# Template builder helpers
# ─────────────────────────────────────────────────────────────────────────────

def _run(para, text, bold=False, size=14, color=None):
    r = para.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(size)
    r.font.bold  = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r


def _para(doc, text='', bold=False, size=14, align=WD_ALIGN_PARAGRAPH.LEFT,
          before=0, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        _run(p, text, bold=bold, size=size)
    return p


def _heading(doc, text, level=1):
    sizes   = {1: 17, 2: 14, 3: 13}
    size    = sizes.get(level, 14)
    align   = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p = _para(doc, text, bold=True, size=size, align=align,
              before=(10 if level == 1 else 6), after=5)
    return p


def _shd(cell, hex_color):
    """Set table cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _cell(cell, text, bold=False, size=13, bg=None, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    _run(p, text, bold=bold, size=size, color=color)
    if bg:
        _shd(cell, bg)


def _cell_para(cell, text, bold=False, size=13, color=None):
    """Add an extra paragraph to an existing cell."""
    p = cell.add_paragraph()
    _run(p, text, bold=bold, size=size, color=color)
    return p


def _sig_table(doc, left_title, left_name_ph, left_pos_ph,
               right_title, right_name_ph, right_pos_ph):
    """Two-column signature block."""
    _para(doc)
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    for c, title in ((0, left_title), (1, right_title)):
        _cell(t.cell(0, c), title, bold=True, size=13,
              bg='1F3864', color=(255, 255, 255),
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in (0, 1):
        _cell(t.cell(1, c), 'ลงชื่อ  ............................................',
              size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(t.cell(2, 0), left_name_ph,  size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(t.cell(2, 1), right_name_ph, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(t.cell(3, 0), left_pos_ph,   size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(t.cell(3, 1), right_pos_ph,  size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in (0, 1):
        _cell(t.cell(4, c), 'วันที่  ........  /  ........  /  ................',
              size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    return t


# ─────────────────────────────────────────────────────────────────────────────
# Build the base .docx template
# ─────────────────────────────────────────────────────────────────────────────

def build_template():
    """Create agreement_template.docx with docxtpl placeholders."""
    doc = Document()

    # ── A4 page ──────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0);  sec.page_height    = Cm(29.7)
    sec.left_margin   = Cm(2.5);   sec.right_margin   = Cm(2.5)
    sec.top_margin    = Cm(2.5);   sec.bottom_margin  = Cm(2.0)

    # Default style
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(14)
    doc.styles['Normal'].paragraph_format.space_after = Pt(4)

    # ── Header: org left | doc number right ──────────────────────────────────
    ht = doc.add_table(rows=1, cols=2)
    ht.style = 'Table Grid'
    ht.columns[0].width = Cm(11)
    ht.columns[1].width = Cm(6)

    lc = ht.cell(0, 0)
    lc.paragraphs[0].clear()
    _run(lc.paragraphs[0], '{{ tester_company_th }}', bold=True, size=14)
    _cell_para(lc, '{{ tester_company_en }}', size=12)
    _cell_para(lc, '{{ tester_address }}',    size=11)

    rc = ht.cell(0, 1)
    rc.paragraphs[0].clear()
    rl = rc.paragraphs[0]
    _run(rl, 'เลขที่เอกสาร:  ', bold=True, size=12)
    _run(rl, '{{ doc_number }}', size=12)
    p_dt = rc.add_paragraph()
    _run(p_dt, 'วันที่:  ', bold=True, size=12)
    _run(p_dt, '{{ doc_date }}', size=12)

    # ── Title ─────────────────────────────────────────────────────────────────
    _para(doc)
    _heading(doc, 'หนังสือขอบเขตการทำ Penetration Testing', level=1)
    _heading(doc, '(Penetration Testing Scope of Work)', level=1)

    # ── 1. คู่สัญญา ──────────────────────────────────────────────────────────
    _heading(doc, '1.  คู่สัญญา', level=2)
    pt = doc.add_table(rows=2, cols=2)
    pt.style = 'Table Grid'
    _cell(pt.cell(0, 0), 'ผู้ว่าจ้าง (Client)',
          bold=True, size=13, bg='2E4057', color=(255, 255, 255),
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(pt.cell(0, 1), 'ผู้ให้บริการ (Service Provider)',
          bold=True, size=13, bg='2E4057', color=(255, 255, 255),
          align=WD_ALIGN_PARAGRAPH.CENTER)

    cc = pt.cell(1, 0)
    cc.paragraphs[0].clear()
    _run(cc.paragraphs[0], 'ชื่อ:  ', bold=True, size=13)
    _run(cc.paragraphs[0], '{{ client_name_th }}', size=13)
    p_en = cc.add_paragraph(); _run(p_en, '{{ client_name_en }}', size=12)
    p_ad = cc.add_paragraph()
    _run(p_ad, 'ที่อยู่:  ', bold=True, size=12)
    _run(p_ad, '{{ client_address }}', size=12)
    p_ct = cc.add_paragraph()
    _run(p_ct, 'ผู้ติดต่อ:  ', bold=True, size=12)
    _run(p_ct, '{{ client_contact }}', size=12)

    tc = pt.cell(1, 1)
    tc.paragraphs[0].clear()
    _run(tc.paragraphs[0], 'ชื่อ:  ', bold=True, size=13)
    _run(tc.paragraphs[0], '{{ tester_company_th }}', size=13)
    p_te = tc.add_paragraph(); _run(p_te, '{{ tester_company_en }}', size=12)
    p_tm = tc.add_paragraph()
    _run(p_tm, 'ทีมงาน:  ', bold=True, size=12)
    _run(p_tm, '{{ team_members_inline }}', size=12)

    # ── 2. ระบบที่ทดสอบ ───────────────────────────────────────────────────────
    _heading(doc, '2.  ระบบที่ทดสอบ (Target Systems)', level=2)
    p = _para(doc, '{{ target_systems }}', size=14)
    p.paragraph_format.left_indent = Cm(0.5)

    # ── 3. รายละเอียดขอบเขต ───────────────────────────────────────────────────
    _heading(doc, '3.  รายละเอียดขอบเขตการทดสอบ', level=2)
    p = _para(doc, '{{ scope_description }}', size=14)
    p.paragraph_format.left_indent = Cm(0.5)

    # ── 4. สิ่งที่ไม่อยู่ในขอบเขต ────────────────────────────────────────────
    _heading(doc, '4.  สิ่งที่ไม่อยู่ในขอบเขต (Out of Scope)', level=2)
    p = _para(doc, '{{ out_of_scope }}', size=14)
    p.paragraph_format.left_indent = Cm(0.5)

    # ── 5. ระยะเวลา ───────────────────────────────────────────────────────────
    _heading(doc, '5.  ระยะเวลาการทดสอบ', level=2)
    tt = doc.add_table(rows=3, cols=2)
    tt.style = 'Table Grid'
    for row, (label, val) in enumerate([
        ('วันเริ่มต้น',       '{{ test_start_date }}'),
        ('วันสิ้นสุด',        '{{ test_end_date }}'),
        ('ช่วงเวลาทำงาน',    '{{ test_hours }}'),
    ]):
        _cell(tt.cell(row, 0), label, bold=True, size=13, bg='F2F2F2')
        _cell(tt.cell(row, 1), val, size=13)

    # ── 6. วิธีการทดสอบ ────────────────────────────────────────────────────────
    _heading(doc, '6.  วิธีการทดสอบ (Methodology)', level=2)
    p = _para(doc, '{{ methodology }}', size=14)
    p.paragraph_format.left_indent = Cm(0.5)

    # ── 7. ทีมงานผู้ทดสอบ ─────────────────────────────────────────────────────
    _heading(doc, '7.  ทีมงานผู้ทดสอบ', level=2)
    p = _para(doc, '{{ team_members }}', size=14)
    p.paragraph_format.left_indent = Cm(0.5)

    # ── Signature: Scope ──────────────────────────────────────────────────────
    _sig_table(
        doc,
        'ผู้ว่าจ้าง', '( {{ client_signer_name }} )', '{{ client_signer_title }}',
        'ผู้ให้บริการ', '( {{ tester_signer_name }} )', '{{ tester_signer_title }}',
    )

    # ═════════════════════════════════════════════════════════════════════════
    # PAGE BREAK → SECTION 2: NDA
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    # Header (same org info)
    ht2 = doc.add_table(rows=1, cols=2)
    ht2.style = 'Table Grid'
    ht2.columns[0].width = Cm(11)
    ht2.columns[1].width = Cm(6)

    lc2 = ht2.cell(0, 0)
    lc2.paragraphs[0].clear()
    _run(lc2.paragraphs[0], '{{ tester_company_th }}', bold=True, size=14)
    _cell_para(lc2, '{{ tester_company_en }}', size=12)

    rc2 = ht2.cell(0, 1)
    rc2.paragraphs[0].clear()
    rl2 = rc2.paragraphs[0]
    _run(rl2, 'เลขที่เอกสาร:  ', bold=True, size=12)
    _run(rl2, '{{ doc_number }}-NDA', size=12)
    p_dt2 = rc2.add_paragraph()
    _run(p_dt2, 'วันที่:  ', bold=True, size=12)
    _run(p_dt2, '{{ doc_date }}', size=12)

    _para(doc)
    _heading(doc, 'ข้อตกลงการรักษาความลับ', level=1)
    _heading(doc, '(Non-Disclosure Agreement)', level=1)

    # Intro
    _para(doc, (
        'ข้อตกลงนี้ทำขึ้นระหว่าง  {{ client_name_th }}  (ต่อไปเรียกว่า "ผู้รับข้อมูล")  '
        'และ  {{ tester_company_th }}  (ต่อไปเรียกว่า "ผู้ให้บริการ")  '
        'โดยทั้งสองฝ่ายตกลงดังต่อไปนี้'
    ), size=14, after=6)

    nda_clauses = [
        ('ข้อ 1  นิยาม',
         '"ข้อมูลที่เป็นความลับ" หมายถึง ข้อมูล รายงาน ผลการทดสอบ ช่องโหว่ที่พบ '
         'รหัสผ่าน แผนผังระบบ และข้อมูลทางเทคนิคใดๆ ที่ได้รับหรือเกิดขึ้นจากการ '
         'ดำเนินการ Penetration Testing ตามหนังสือขอบเขตฉบับนี้'),
        ('ข้อ 2  พันธกรณีในการรักษาความลับ',
         'ผู้ให้บริการตกลงที่จะ (ก) เก็บรักษาข้อมูลที่เป็นความลับไว้เป็นความลับ  '
         '(ข) ไม่เปิดเผยต่อบุคคลภายนอก  (ค) ใช้ข้อมูลเพื่อวัตถุประสงค์ตามสัญญานี้เท่านั้น  '
         '(ง) จำกัดการเข้าถึงเฉพาะทีมงานที่เกี่ยวข้องโดยตรง'),
        ('ข้อ 3  ระยะเวลา',
         'ข้อตกลงนี้มีผลบังคับใช้นับตั้งแต่วันที่ลงนาม และมีผลต่อเนื่องเป็นระยะเวลา  '
         '{{ nda_duration_years }}  ปี  แม้ว่างานตามสัญญาหลักจะสิ้นสุดลงแล้วก็ตาม'),
        ('ข้อ 4  ข้อยกเว้น',
         'พันธกรณีข้างต้นไม่มีผลกับข้อมูลที่ (ก) เป็นสาธารณะโดยไม่ใช่ความผิดของผู้ให้บริการ  '
         '(ข) ผู้ให้บริการได้รับจากบุคคลที่สามอย่างถูกต้องชอบธรรม  '
         '(ค) ต้องเปิดเผยตามคำสั่งศาลหรือกฎหมาย'),
        ('ข้อ 5  ผลของการละเมิด',
         'การละเมิดข้อตกลงนี้จะก่อให้เกิดความเสียหายที่ไม่สามารถประเมินได้เป็นตัวเงิน '
         'ผู้รับข้อมูลมีสิทธิ์ขอคำสั่งห้ามจากศาล นอกเหนือจากสิทธิ์เรียกค่าเสียหายตามกฎหมาย'),
        ('ข้อ 6  กฎหมายที่ใช้บังคับ',
         'ข้อตกลงนี้อยู่ภายใต้บังคับและตีความตามกฎหมายแห่งราชอาณาจักรไทย '
         'และให้ศาลไทยเป็นผู้มีเขตอำนาจพิจารณาข้อพิพาท'),
    ]

    for title, body in nda_clauses:
        _heading(doc, title, level=2)
        p = _para(doc, body, size=14, after=6)
        p.paragraph_format.left_indent       = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(0.5)

    # ── Signature: NDA ────────────────────────────────────────────────────────
    _sig_table(
        doc,
        'ผู้รับข้อมูล', '( {{ client_signer_name }} )', '{{ client_signer_title }}',
        'ผู้ให้บริการ', '( {{ tester_signer_name }} )', '{{ tester_signer_title }}',
    )

    doc.save(TEMPLATE_PATH)
    logger.info('Agreement template created: %s', TEMPLATE_PATH)
    return TEMPLATE_PATH


# ─────────────────────────────────────────────────────────────────────────────
# Render + export
# ─────────────────────────────────────────────────────────────────────────────

def _context(agreement, org_id=None):
    """Build the Jinja2 context dict from a PentestAgreement."""
    from scanner.models import OrganizationProfile
    import datetime

    org = OrganizationProfile.load(org_id=org_id or agreement.org_id)

    def fmt_date(d):
        if not d:
            return '—'
        # Buddhist year
        return d.strftime(f'%d/%m/') + str(d.year + 543)

    team_lines = [l.strip() for l in (agreement.team_members or '').splitlines() if l.strip()]
    team_inline = ',  '.join(team_lines) if team_lines else '—'

    return {
        # Org (tester side)
        'tester_company_th':  org.name_th or '—',
        'tester_company_en':  org.name_en or '',
        'tester_address':     org.address or '',
        'tester_signer_name':  agreement.tester_signer_name or org.approver_name or '—',
        'tester_signer_title': agreement.tester_signer_title or org.approver_title or '',
        # Document
        'doc_number': agreement.document_number or '—',
        'doc_date':   fmt_date(agreement.created_at.date() if agreement.created_at else None),
        # Client
        'client_name_th':      agreement.client_name_th or '—',
        'client_name_en':      agreement.client_name_en or '',
        'client_address':      agreement.client_address or '—',
        'client_contact':      agreement.client_contact or '—',
        'client_signer_name':  agreement.client_signer_name or '—',
        'client_signer_title': agreement.client_signer_title or '',
        # Scope
        'target_systems':      agreement.target_systems   or '—',
        'scope_description':   agreement.scope_description or '—',
        'out_of_scope':        agreement.out_of_scope     or '—',
        'methodology':         agreement.methodology      or '—',
        'team_members':        agreement.team_members     or '—',
        'team_members_inline': team_inline,
        # Timeline
        'test_start_date': fmt_date(agreement.test_start_date),
        'test_end_date':   fmt_date(agreement.test_end_date),
        'test_hours':      agreement.test_hours or '09:00 – 18:00',
        # NDA
        'nda_duration_years': str(agreement.nda_duration_years or 3),
    }


def generate_agreement_docx(agreement, org_id=None):
    """Render template → return .docx bytes."""
    from docxtpl import DocxTemplate

    if not os.path.exists(TEMPLATE_PATH):
        build_template()

    tpl = DocxTemplate(TEMPLATE_PATH)
    tpl.render(_context(agreement, org_id))

    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def generate_agreement_pdf_docx(agreement, org_id=None):
    """Render template → convert to PDF via LibreOffice headless → return bytes."""
    docx_bytes = generate_agreement_docx(agreement, org_id)

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, 'agreement.docx')
        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [
                'soffice', '--headless', '--invisible',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', tmp,
                docx_path,
            ],
            capture_output=True,
            timeout=60,
            env={**os.environ, 'HOME': tmp},   # isolate LO user profile
        )
        if result.returncode != 0:
            raise RuntimeError(
                f'LibreOffice conversion failed: {result.stderr.decode(errors="replace")}'
            )

        pdf_path = os.path.join(tmp, 'agreement.pdf')
        with open(pdf_path, 'rb') as f:
            return f.read()
